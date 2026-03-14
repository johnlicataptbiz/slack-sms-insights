from __future__ import annotations

from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUT_PATH = "/Users/jl/Desktop/SlackCLI/output/doc/ptbizsms_audit_future_proofing_2026-02-22.docx"


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sp = doc.add_paragraph(subtitle)
    sp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sp.runs[0].italic = True

    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code.strip("\n"))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)


def priority_line(doc: Document, rating: str) -> None:
    p = doc.add_paragraph(f"Priority: {rating}")
    p.runs[0].bold = True


def build() -> None:
    doc = Document()
    set_page(doc)
    set_styles(doc)

    add_title(
        doc,
        "PTBizSMS Audit and Future-Proofing Roadmap",
        "Production URL: https://ptbizsms.com | Audit date: February 22, 2026",
    )

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    add_bullets(
        doc,
        [
            "Overall health score: D (42.9/100 weighted) as of February 22, 2026. Weighted breakdown: Security 40% (20/100), Reliability 20% (63/100), Maintainability 15% (55/100), UX 15% (64/100), Scalability/Cost 10% (45/100). The product is useful and actively used, but current auth/data-exposure posture is dangerously below production SaaS standards.",
            "3 biggest wins: (1) V2 dashboard is materially better than legacy in clarity and workflow depth (/v2/insights, /v2/inbox, /v2/runs), (2) backend has strong operational primitives (idempotency + send-attempt audit trail in /Users/jl/Desktop/SlackCLI/sms-insights/services/db.ts:332), (3) reporting/attribution logic is rich and differentiating for coaching teams.",
            "3 biggest risks: (1) sensitive read/write endpoints are publicly reachable without auth (/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:2053, /Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:2066, /Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:2079), (2) OAuth token is written to localStorage, increasing theft risk (/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:297, /Users/jl/Desktop/SlackCLI/frontend/src/api/client.ts:42), (3) no send-rate guardrails on public mutation paths creates direct SMS cost-explosion risk.",
            "Bottom line: This is a solid product trapped behind critical trust-boundary flaws. Fixing auth/session/abuse controls in the next 30 days will unlock safer scale without re-platforming.",
        ],
    )
    priority_line(doc, "1/5")

    # 2. Live Site Crawl Report
    doc.add_heading("2. Live Site Crawl Report", level=1)
    doc.add_heading("Full sitemap discovered", level=2)
    add_table(
        doc,
        ["Flow Group", "Route", "Result", "Notes"],
        [
            ["Entry/Auth", "https://ptbizsms.com/", "200", "Password gate shown first."],
            ["Legacy", "/legacy", "200", "Legacy shell with Team Insights, Daily Runs, rep scorecards, deep dives."],
            ["Legacy", "/legacy/*", "200", "Same legacy app shell route handling."],
            ["V2", "/v2/insights", "200", "KPI-heavy overview + weekly summary panel."],
            ["V2", "/v2/inbox", "200", "Conversation list/detail, draft generation, send, qualification/escalation controls."],
            ["V2", "/v2/runs", "200", "Run list + detail + raw report expansion + share links."],
            ["V2", "/v2/rep/jack", "200", "Jack scorecard."],
            ["V2", "/v2/rep/brandon", "200", "Brandon scorecard."],
            ["V2", "/v2/sequences", "200", "Sequence analytics and audit rows."],
            ["V2 alias", "/v2/attribution", "Redirect", "Redirects to /v2/sequences."],
            ["V2 invalid", "/v2/unknown-route", "200 view-level not-found", "Route not found in v2 layer."],
            ["API (observed)", "/api/v2/runs?limit=1", "200", "Returns full run payload incl. full report text."],
            ["API (observed)", "/api/v2/inbox/conversations?limit=1", "200", "Returns names/phones/message snippets without auth token."],
            ["API (observed)", "/api/v2/inbox/conversations/:id", "200", "Returns full thread/contact card without auth token."],
            ["API (observed)", "/api/auth/verify", "401", "One of few endpoints that requires auth."],
        ],
    )

    doc.add_heading("Screenshot-style description of major screens", level=2)
    add_table(
        doc,
        ["Screen", "Visual/Layout description", "Interaction behavior"],
        [
            ["Password Gate", "Centered card, PT Biz logo, single password input, Stay logged in checkbox.", "Client-side check only; unlock state via cookie/session storage."],
            ["Legacy Shell", "Dense left nav with emoji labels; content area packed with metrics/tables.", "Functional but visually dated; high cognitive load."],
            ["Team Insights (Legacy)", "Large KPI blocks and report-derived text sections.", "Heavy text scanning required; weak hierarchy for what to do next."],
            ["Daily Runs (Legacy)", "Accordion-like run history with expandable raw report details.", "Useful forensic view; includes sensitive full report text in-browser."],
            ["V2 Insights", "Cleaner card layout, KPI grid, weekly summary panel, source badges.", "Best visual polish in app; still dense under weekly table sections."],
            ["V2 Inbox", "3-column feel (filters/list/detail), composer modal, send-line controls.", "Strong workflow depth; no explicit guardrails in UI for unauthorized backend calls."],
            ["V2 Runs", "Filter bar + selected run detail + parsed metrics + raw report expansion.", "Good operational utility; still exposes complete raw report payload."],
            ["V2 Sequences", "Table-driven sequence performance with audit details.", "Useful for coaching diagnosis; attribution terminology can confuse non-analysts."],
        ],
    )

    doc.add_heading("UX flow analysis", level=2)
    add_bullets(
        doc,
        [
            "Onboarding/auth flow: Entering password opens app immediately, but this is not real access control. It behaves like a cosmetic gate rather than security.",
            "Campaign/reply flow: There is no full campaign builder UI in this app; outbound action is centered in Inbox per-conversation send. Draft-and-send flow is fast but lacks visible compliance/approval checkpoints.",
            "Reporting flow: Insights -> Runs -> Sequence drill-down is coherent and useful for coaches/operators; strongest product area.",
            "Settings flow: No dedicated settings area discovered. Send-line default is embedded inside Inbox; discoverability is low.",
            "Dead ends/secondary auth: No secondary login discovered beyond password gate. /v2/attribution is a redirect alias, not a distinct module.",
        ],
    )

    doc.add_heading("Mobile vs desktop discrepancies", level=2)
    add_table(
        doc,
        ["Area", "Desktop", "Mobile"],
        [
            ["Navigation", "Sidebar + top actions works well.", "Hamburger + overlay works, but deep navigation feels cramped."],
            ["KPI Definitions drawer", "Usable side drawer.", "In one deep-scroll state, close interaction became unreliable due viewport/overlay positioning."],
            ["Data tables", "Dense but readable on wide screens.", "Horizontal density and long rows degrade scan speed significantly."],
            ["Legacy pages", "Already dense.", "Becomes very crowded and operationally tiring."],
        ],
    )

    doc.add_heading("Broken links, 404s, console errors, confusing copy", level=2)
    add_bullets(
        doc,
        [
            "No broken in-app nav links found in sampled flows.",
            "/v2/unknown-route correctly renders internal not-found.",
            "Runtime anomaly: non-UUID conversation path to send endpoint returned 500 instead of 400 (repro with /api/v2/inbox/conversations/does-not-exist/send).",
            "Confusing copy area: Calls booked (Slack) vs SMS booking hints diagnostic is accurate but easy for operators to misread without stronger hierarchy.",
        ],
    )
    priority_line(doc, "2/5")

    # 3. Tech stack
    doc.add_heading("3. Tech Stack & Architecture Deep Dive", level=1)
    add_table(
        doc,
        ["Layer", "Current implementation", "Assessment"],
        [
            ["Frontend", "React 19 + TypeScript + Vite + React Router + TanStack Query (/Users/jl/Desktop/SlackCLI/frontend/package.json:12)", "Modern stack; good base."],
            ["Route architecture", "App-level split between legacy and v2 (/Users/jl/Desktop/SlackCLI/frontend/src/App.tsx:78, /Users/jl/Desktop/SlackCLI/frontend/src/App.tsx:80)", "Transitional architecture causes duplication/drift risk."],
            ["Backend", "Node + TypeScript + Slack Bolt + custom http server (/Users/jl/Desktop/SlackCLI/sms-insights/app.ts:30, /Users/jl/Desktop/SlackCLI/sms-insights/app.ts:66)", "Works, but monolithic API routing and weak policy enforcement."],
            ["API routing", "Single large route registry with requiresAuth per-route (/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:2046)", "Biggest security/process risk zone."],
            ["Database", "PostgreSQL via pg; schema initialized at runtime (/Users/jl/Desktop/SlackCLI/sms-insights/services/db.ts:57)", "Additive, pragmatic, but lacks formal migrations/version discipline."],
            ["SMS integration", "Aloware API wrappers (/Users/jl/Desktop/SlackCLI/sms-insights/services/aloware-client.ts:203) + send orchestration (/Users/jl/Desktop/SlackCLI/sms-insights/services/inbox-send.ts:102)", "Good abstraction start; needs stronger auth/abuse gating around caller paths."],
            ["Auth/session", "Slack OAuth callback writes token to browser localStorage (/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:297)", "High-risk model for production web app."],
            ["Deployment", "Vercel static frontend + rewrite to Railway API (/Users/jl/Desktop/SlackCLI/vercel.json:7)", "Simple and workable; split observability/security controls needed."],
            ["Monitoring", "App logs + Slack error-posting (/Users/jl/Desktop/SlackCLI/sms-insights/services/error-reporter.ts:19)", "Basic; lacks APM, SLO dashboards, anomaly alerting for spend/abuse."],
        ],
    )
    priority_line(doc, "2/5")

    # 4. Critical issues
    doc.add_heading("4. Critical Issues", level=1)
    add_table(
        doc,
        ["Severity", "Issue", "File/Location", "Impact", "Recommended Fix"],
        [
            ["Critical", "OWASP A01 Broken Access Control: password gate is client-side only", "/Users/jl/Desktop/SlackCLI/frontend/src/components/PasswordGate.tsx:5", "Anyone can bypass by editing storage/cookies or reading bundle.", "Remove as security control; enforce server session auth on all dashboard/API routes."],
            ["Critical", "OWASP A01: sensitive read endpoints publicly exposed (runs, inbox)", "/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:2053", "Unauthenticated users can pull reports, conversations, phone numbers.", "Default-deny auth policy; make exceptions explicit and minimal."],
            ["Critical", "OWASP A01/A04: mutating inbox endpoints are public (send, qualification, escalation)", "/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:2079", "Unauthorized state changes and potential SMS sends.", "Require auth + role checks + tenant checks + CSRF protection."],
            ["Critical", "OWASP A07 Identification & Authentication Failures: OAuth token stored in localStorage", "/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:297", "Token theft via XSS/extensions/session replay.", "Use server session + HttpOnly cookie; remove token write script."],
            ["Critical", "OWASP A05 Security Misconfiguration: dummy auth bypass token supported", "/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:139", "If env is mis-set, auth bypass is trivial.", "Remove feature in production builds; fail startup if enabled outside local."],
            ["High", "No rate limits on expensive/public mutation paths", "/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:2109", "SMS spend explosion and API abuse.", "Add rate limiting at edge + app (IP, user, conversation, tenant) with 429 + retry headers."],
            ["High", "Shared static secret (x-bot-token) for /api/runs ingestion", "/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:358", "Replay/secret leakage risk; weak origin assurance.", "Replace with HMAC-signed request (timestamp + nonce + signature) and short replay window."],
            ["High", "PII-rich contracts shipped broadly (contactPhone, full report text)", "/Users/jl/Desktop/SlackCLI/sms-insights/api/v2-contract.ts:135", "Unnecessary data exposure increases legal and breach blast radius.", "Introduce response redaction tiers by role and endpoint purpose."],
            ["High", "Tenant boundary absent in core entities", "/Users/jl/Desktop/SlackCLI/sms-insights/services/db.ts:152", "Unsafe to scale multi-practice without cross-account data risk.", "Add practice_id to all core tables and enforce in every query/index."],
            ["Medium", "Invalid IDs trigger 500 instead of 400", "/Users/jl/Desktop/SlackCLI/sms-insights/services/conversation-store.ts:22", "Noisy ops and easier endpoint probing; poor API hygiene.", "Validate UUID route params before DB call and return typed 400 errors."],
            ["Medium", "JSON body parser has no payload size cap", "/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:103", "Memory pressure / DoS vector with oversized bodies.", "Enforce max body size + 413 Payload Too Large."],
            ["Medium", "Missing explicit CSP/frameguard in app-layer responses", "/Users/jl/Desktop/SlackCLI/sms-insights/app.ts:66", "Weaker browser-side hardening against XSS/clickjacking.", "Set CSP, frame-ancestors, referrer policy at Vercel/edge and backend."],
        ],
    )

    doc.add_heading("Live runtime proof points (Feb 22, 2026)", level=2)
    add_code_block(
        doc,
        """
GET /api/v2/runs?limit=1                     -> 200 with fullReport payload
GET /api/v2/inbox/conversations?limit=1      -> 200 with contactPhone/message data
GET /api/v2/inbox/conversations/{id}         -> 200 with full thread/contact card
POST /api/v2/inbox/conversations/{uuid}/send -> 404 Conversation not found (endpoint reached without auth)
GET /api/auth/verify                         -> 401 (auth actually enforced here)
""",
    )
    priority_line(doc, "1/5")

    # 5. Code quality
    doc.add_heading("5. Code Quality & Maintainability", level=1)
    doc.add_heading("Architecture smells", level=2)
    add_numbered(
        doc,
        [
            "Route policy drift and contradiction: comment says Authenticated endpoints, but many are requiresAuth: false (/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:2061).",
            "Monolithic API file: /Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts mixes auth, CORS, parsing, metrics, inbox, runs, and business logic.",
            "Runtime schema migration pattern: /Users/jl/Desktop/SlackCLI/sms-insights/services/db.ts:57 applies schema at startup, brittle with parallel deploys.",
            "Legacy + v2 dual UI surface increases copy drift and support overhead.",
        ],
    )

    doc.add_heading("Duplication, tight coupling, missing tests", level=2)
    add_bullets(
        doc,
        [
            "No dedicated API route security tests under /Users/jl/Desktop/SlackCLI/sms-insights/tests for requiresAuth policy.",
            "Frontend and backend carry auth/session assumptions, but no contract tests enforce them.",
            "Ingestion/reporting/inbox concerns are coupled in route handlers instead of feature modules.",
        ],
    )

    doc.add_heading("Performance anti-patterns", level=2)
    add_bullets(
        doc,
        [
            "N-request trend hydration: Promise.all(uniqueDays.map(fetchV2SalesMetrics)) in /Users/jl/Desktop/SlackCLI/frontend/src/api/v2Queries.ts:117.",
            "Heavy payload on list endpoints: runs list includes fullReport on every row (/Users/jl/Desktop/SlackCLI/sms-insights/api/v2-contract.ts:561).",
            "Large unpaginated detail reads: inbox detail fetches up to 250 messages per open (/Users/jl/Desktop/SlackCLI/sms-insights/api/routes.ts:1494).",
        ],
    )

    doc.add_heading("Specific refactors with before/after snippets", level=2)
    doc.add_paragraph("Refactor A - Default-deny auth route policy")
    add_code_block(
        doc,
        """
// Before
{ method: 'GET', path: '/api/v2/inbox/conversations', requiresAuth: false, handler: ... }
{ method: 'POST', path: '/api/v2/inbox/conversations/:id/send', requiresAuth: false, handler: ... }

// After
const publicRoutes = new Set([
  'GET /api/oauth/start',
  'GET /api/oauth/callback',
  'GET /api/health',
]);
const requiresAuth = (method: string, path: string) => !publicRoutes.has(`${method} ${path}`);
""",
    )

    doc.add_paragraph("Refactor B - Replace localStorage token with secure session cookie")
    add_code_block(
        doc,
        """
// Before
localStorage.setItem('slackToken', token);
const t = localStorage.getItem('slackToken');
if (t) headers.set('Authorization', `Bearer ${t}`);

// After
Set-Cookie: session=<opaque-id>; HttpOnly; Secure; SameSite=Lax; Path=/;
fetch(path, { credentials: 'include', ...init });
""",
    )

    doc.add_paragraph("Refactor C - Batch trend endpoint to remove N+1 front-end requests")
    add_code_block(
        doc,
        """
// Before
const envelopes = await Promise.all(uniqueDays.map((day) => fetchV2SalesMetrics({ day, tz })));

// After
const response = await client.get(`/api/v2/sales-metrics/batch?${params.toString()}`);
""",
    )
    priority_line(doc, "2/5")

    # 6. UI/UX
    doc.add_heading("6. UI/UX & Conversion Opportunities", level=1)
    add_table(
        doc,
        ["Priority", "Opportunity", "Why it matters for PT owners", "Exact implementation suggestion"],
        [
            ["1", "Today Action Bar in Insights", "Owners need immediate actions, not just diagnostics.", "Add top strip: Needs Reply, At-Risk Sequences, Booked Today; each links to filtered Inbox/Sequences."],
            ["1", "Inbox Next Best Action chip", "Reduces decision fatigue during follow-up blocks.", "In /Users/jl/Desktop/SlackCLI/frontend/src/v2/pages/InboxV2.tsx, derive chip from escalation + qualification state."],
            ["2", "Stronger metric hierarchy", "Current KPI names can be misread by non-analysts.", "Keep canonical KPI first, move diagnostic KPIs into collapsible Diagnostics panel."],
            ["2", "Premium visual system pass", "Current UI is functional but not premium-coach quality.", "Introduce tokenized visual variables in /Users/jl/Desktop/SlackCLI/frontend/src/v2/v2.css."],
            ["2", "Mobile-first table handling", "Mobile scan cost is high on dense tables.", "Collapse low-priority columns into expandable row details at <768px."],
            ["3", "Consistent empty/loading/error states", "Confidence drops when states are abrupt.", "Standardize loading/empty/error/retry card component across v2 pages."],
            ["3", "Dedicated Settings page", "Send defaults are buried in Inbox flow.", "Add /v2/settings for outbound defaults, notifications, role visibility."],
            ["3", "Conversion-friendly copy updates", "Current copy is accurate but not outcome-oriented.", "Use owner language: Replies awaiting action, At-risk follow-up, Booked-call confidence."],
        ],
    )

    doc.add_heading("Concrete CSS direction", level=2)
    add_code_block(
        doc,
        """
:root {
  --surface: #ffffff;
  --surface-elevated: #f8fafc;
  --ink: #0f172a;
  --muted: #475569;
  --accent: #0f766e;
  --danger: #b91c1c;
  --radius-lg: 16px;
  --shadow-soft: 0 8px 24px rgba(15, 23, 42, 0.08);
}

.V2MetricCard {
  border-radius: var(--radius-lg);
  box-shadow: var(--shadow-soft);
  border: 1px solid #e2e8f0;
  background: linear-gradient(180deg, var(--surface) 0%, var(--surface-elevated) 100%);
}
""",
    )
    priority_line(doc, "3/5")

    # 7. Features
    doc.add_heading("7. Feature & Product Ideas", level=1)
    doc.add_heading("Quick Wins (1-2 weeks)", level=2)
    add_table(
        doc,
        ["Idea", "Build shape", "Business impact"],
        [
            ["SLA Breach Digest", "Daily 7am digest: conversations > X min unreplied.", "Immediate follow-up improvement, fewer lost leads."],
            ["Coach-approved reply templates", "Curated snippets by objection type with lint + edit requirement.", "Faster reps, more consistent brand voice."],
            ["Inbox intent auto-tags", "Tag inbound messages (pricing, schedule, not interested) with confidence.", "Better triage and sequence optimization loops."],
            ["Last 24h anomalies panel", "Flag opt-out spike, reply-rate drop, send failure spike.", "Earlier intervention before revenue impact."],
        ],
    )

    doc.add_heading("2-4 week ideas", level=2)
    add_table(
        doc,
        ["Idea", "Build shape", "Business impact"],
        [
            ["Sequence Attribution Debugger", "Per-booked-call explainability view from audit rows.", "Trust in attribution and coaching recommendations."],
            ["Multi-location rollup", "Add location dimension + comparison cards.", "Better franchise/group visibility."],
            ["Rep coaching score", "Composite score from response lag, conversion, opt-outs.", "Better manager accountability and coaching cadence."],
            ["AI inbox copilot with approval policy", "Draft + rationale + must-edit conditions for risky sends.", "Higher productivity without losing control."],
        ],
    )

    doc.add_heading("Strategic ideas (6+ months)", level=2)
    add_table(
        doc,
        ["Idea", "Build shape", "Business impact"],
        [
            ["True multi-tenant white-label portal", "practice_id tenancy, RBAC, branded domains, per-tenant policies.", "New revenue channel and safer scale."],
            ["Spend optimizer", "Predictive send pacing and budget guardrails by campaign/tenant.", "Direct margin protection at 100x SMS volume."],
            ["Patient lifecycle intelligence", "Merge SMS events + booked call outcomes + pipeline stages into forecasting model.", "Stronger retention and upsell value."],
            ["Automation marketplace", "Reusable automations for reminders/follow-up/reactivation with compliance controls.", "Product stickiness and differentiation."],
        ],
    )
    priority_line(doc, "3/5")

    # 8. Performance roadmap
    doc.add_heading("8. Performance, Scalability & Cost Roadmap", level=1)
    doc.add_heading("10x practices / 100x SMS target design", level=2)
    add_numbered(
        doc,
        [
            "Move outbound sends to queue-first architecture: API writes send intent row, workers consume queue with provider throttles/retries, preserving idempotency.",
            "Introduce hard spend controls: per-practice/global hourly/day/week caps, anomaly detection, and auto-pausing rules with manual override.",
            "Add strict abuse controls: rate limits by IP/user/practice, CSRF for cookie-auth mutations, replay-safe signatures for machine ingestion endpoints.",
            "Reduce payload and compute load: summary-only run list by default, batch metrics endpoints, cache expensive analytics windows.",
            "Scale data model: add practice_id to operational tables; partition high-volume event tables by date + practice.",
            "Upgrade observability: SLOs, structured logs with correlation IDs, and alerting for auth failures/send anomalies.",
        ],
    )

    doc.add_heading("Cost-control guardrails for SMS", level=2)
    add_table(
        doc,
        ["Guardrail", "Trigger", "Action"],
        [
            ["Daily spend cap", ">100% expected daily budget", "Block non-critical sends; alert owner/coaches."],
            ["Burst cap", "5-minute send rate > threshold", "Queue slowdown + notify on-call."],
            ["Opt-out risk guard", "Sequence opt-out > threshold", "Auto-throttle sequence; require explicit re-enable."],
            ["Failed-send spike", "Provider failure ratio > threshold", "Fail over provider route and pause risky batches."],
        ],
    )
    priority_line(doc, "1/5")

    # 9. Implementation roadmap
    doc.add_heading("9. Implementation Roadmap", level=1)
    doc.add_heading("30-day backlog (security stabilization first)", level=2)
    add_table(
        doc,
        ["Workstream", "Key deliverables", "Owner", "Effort", "Impact"],
        [
            ["Auth hardening", "Lock down all sensitive /api/v2/* read/write routes; default-deny policy", "Backend", "M", "Very High"],
            ["Session redesign", "Replace localStorage token model with secure session cookies", "Backend + Frontend", "M", "Very High"],
            ["Abuse controls", "Add rate limits + SMS send caps + idempotency checks", "Backend + Platform", "M", "Very High"],
            ["Input hygiene", "UUID validation + request size limits + structured errors", "Backend", "S", "High"],
            ["Incident visibility", "Add auth/send anomaly alerting and basic SLO dashboard", "Platform", "S", "High"],
        ],
    )

    doc.add_heading("90-day backlog (platform and product quality)", level=2)
    add_table(
        doc,
        ["Workstream", "Key deliverables", "Owner", "Effort", "Impact"],
        [
            ["Multi-tenant foundation", "Add practice_id model and query enforcement", "Backend + DB", "L", "Very High"],
            ["API modularization", "Split monolithic routes by domain + shared middleware", "Backend", "M", "High"],
            ["UX premium pass", "Action-bar, settings page, mobile table patterns", "Frontend + Product", "M", "High"],
            ["Analytics performance", "Batch endpoints + payload slimming + caching", "Backend + Frontend", "M", "High"],
            ["Regression coverage", "Route auth tests + integration tests for critical flows", "Backend QA", "M", "High"],
        ],
    )

    doc.add_heading("6-month backlog (strategic scale and differentiation)", level=2)
    add_table(
        doc,
        ["Workstream", "Key deliverables", "Owner", "Effort", "Impact"],
        [
            ["White-label portal", "Tenant branding, RBAC, domain strategy", "Product + Eng", "L", "Very High"],
            ["Automation intelligence", "AI-assisted reply/workflow orchestration with governance", "Product + Eng", "L", "High"],
            ["Cost optimizer", "Predictive spend controls and send pacing engine", "Data + Platform", "L", "High"],
            ["Enterprise reliability", "DR playbooks, queue failover, provider redundancy", "Platform", "M", "High"],
        ],
    )

    doc.add_heading("Effort vs impact matrix", level=2)
    add_table(
        doc,
        ["Item", "Effort", "Impact", "Sequence"],
        [
            ["Lock down public sensitive endpoints", "M", "Very High", "Week 1"],
            ["Secure session cookies", "M", "Very High", "Week 1-2"],
            ["Rate limiting + send caps", "M", "Very High", "Week 2"],
            ["UUID/body validation", "S", "High", "Week 2"],
            ["Multi-tenant schema foundation", "L", "Very High", "Month 2-3"],
            ["API modularization", "M", "High", "Month 2"],
            ["UX premium pass", "M", "High", "Month 2-3"],
        ],
    )

    doc.add_heading("Required test cases and acceptance scenarios", level=2)
    add_numbered(
        doc,
        [
            "Unauthorized requests to all sensitive GET/POST routes return 401/403.",
            "Cross-tenant access attempts (IDOR) fail across runs/inbox/send state.",
            "Session theft simulation confirms no token exposure in browser storage.",
            "Rate-limit tests enforce 429 and preserve normal traffic quality.",
            "SMS send guardrails block budget overrun and produce auditable events.",
            "Desktop/mobile regression suite passes key flows (insights, inbox, runs).",
            "Rollout safety validates feature flags, staged release, and rollback behavior.",
        ],
    )
    priority_line(doc, "1/5")

    doc.add_paragraph()
    fp = doc.add_paragraph(f"Generated on {date.today().isoformat()} in /Users/jl/Desktop/SlackCLI/output/doc")
    fp.runs[0].italic = True

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
